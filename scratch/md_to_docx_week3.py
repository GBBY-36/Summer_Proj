import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_markdown_formatting(paragraph, text):
    # Match bold (**text**) and code (`text`)
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            cleaned = part[2:-2]
            run = paragraph.add_run(cleaned)
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            cleaned = part[1:-1]
            run = paragraph.add_run(cleaned)
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            # Match markdown link [text](url) -> keep text or text (url)
            sub_parts = re.split(r'(\[.*?\]\(.*?\))', part)
            for sp in sub_parts:
                match = re.match(r'\[(.*?)\]\((.*?)\)', sp)
                if match:
                    link_text = match.group(1)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(46, 117, 182)
                    run.underline = True
                else:
                    paragraph.add_run(sp)

def parse_and_add_table(doc, rows):
    if len(rows) < 3:
        return
    header_row = [cell.strip() for cell in rows[0].split("|")[1:-1]]
    data_rows = []
    for r in rows[2:]:
        data_rows.append([cell.strip() for cell in r.split("|")[1:-1]])
        
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(header_row))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for idx, name in enumerate(header_row):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1F4E79")
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = 'Arial'
            
    for r_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F2F2F2" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            p = row_cells[c_idx].paragraphs[0]
            parse_markdown_formatting(p, val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)

def main():
    md_path = "/Users/qinuoshi/.gemini/antigravity/brain/26a3f220-54ed-4fac-82ed-f1aa3a44c31a/project_description_week3.md"
    docx_path = "/Users/qinuoshi/Desktop/Graduate/PhD/Week3/project_description_week3.docx"

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    in_code_block = False
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if line.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
            continue
            
        if line.startswith("|") and line.endswith("|"):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            parse_and_add_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        if not line.strip():
            i += 1
            continue
            
        # Headings
        if line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(46, 117, 182)
        elif line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(90, 90, 90)
            
        # Lists
        elif line.strip().startswith("* ") or line.strip().startswith("- "):
            indent_level = len(line) - len(line.lstrip())
            text = re.sub(r'^[\s]*[\*\-]\s*', '', line)
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (indent_level // 2))
            parse_markdown_formatting(p, text)
            
        # Blockquote
        elif line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(6)
            
            if text.startswith("[!NOTE]"):
                run_prefix = p.add_run("【注意】 ")
                run_prefix.bold = True
                run_prefix.font.color.rgb = RGBColor(46, 117, 182)
                text = text[7:].strip()
            elif text.startswith("[!IMPORTANT]"):
                run_prefix = p.add_run("【重要】 ")
                run_prefix.bold = True
                run_prefix.font.color.rgb = RGBColor(180, 0, 0)
                text = text[12:].strip()
                
            parse_markdown_formatting(p, text)
            
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            parse_markdown_formatting(p, line)
            
        i += 1
        
    if in_table and table_rows:
        parse_and_add_table(doc, table_rows)
        
    doc.save(docx_path)
    print("Success: saved project_description_week3.docx to", docx_path)

if __name__ == "__main__":
    main()
